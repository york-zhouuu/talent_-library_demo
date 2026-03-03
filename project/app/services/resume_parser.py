import os
import json
import re
import base64
import aiofiles
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from app.core import get_settings
from app.services.ai_service import AIService

settings = get_settings()


class ResumeParser:
    def __init__(self):
        self.ai = AIService()
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def save_file(self, filename: str, content: bytes) -> str:
        # 只取文件名，去掉路径部分（处理文件夹上传的情况）
        safe_filename = Path(filename).name
        file_path = self.upload_dir / safe_filename

        # 处理文件名冲突：如果文件已存在，添加时间戳
        if file_path.exists():
            import time
            stem = file_path.stem
            suffix = file_path.suffix
            safe_filename = f"{stem}_{int(time.time())}{suffix}"
            file_path = self.upload_dir / safe_filename

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)
        return str(file_path)

    def extract_text_from_pdf(self, content: bytes) -> str:
        """从 PDF 提取文本"""
        try:
            reader = PdfReader(BytesIO(content))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            text = text.strip()
            return text
        except Exception as e:
            print(f"PDF 提取失败: {e}")
            return ""

    def is_image_based_pdf(self, content: bytes) -> bool:
        """检测 PDF 是否是图片版（没有可提取的文本）"""
        try:
            reader = PdfReader(BytesIO(content))
            for page in reader.pages:
                # 检查是否有字体资源
                resources = page.get("/Resources", {})
                if resources and "/Font" in resources:
                    return False
            return True
        except:
            return False

    async def extract_text_from_image_pdf(self, content: bytes) -> str:
        """使用 Claude Vision 从图片版 PDF 提取文本"""
        try:
            # 尝试导入 pdf2image
            try:
                from pdf2image import convert_from_bytes
            except ImportError:
                print("pdf2image 未安装，无法处理图片版 PDF")
                return ""

            # 将 PDF 转换为图片
            images = convert_from_bytes(content, dpi=150, first_page=1, last_page=3)  # 只处理前3页

            if not images:
                return ""

            # 将图片转为 base64
            image_contents = []
            for i, img in enumerate(images[:3]):  # 最多3页
                buffer = BytesIO()
                img.save(buffer, format="JPEG", quality=85)
                img_base64 = base64.b64encode(buffer.getvalue()).decode()
                image_contents.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/jpeg",
                        "data": img_base64
                    }
                })

            # 使用 Claude Vision 提取文本
            image_contents.append({
                "type": "text",
                "text": "请仔细阅读这份简历图片，提取并返回简历中的所有文本内容。保持原有格式，包括姓名、联系方式、工作经历、教育背景、技能等信息。只返回提取的文本，不要添加任何解释。"
            })

            response = await self.ai.client.messages.create(
                model=self.ai.model,
                max_tokens=4096,
                messages=[{"role": "user", "content": image_contents}]
            )

            return self.ai._extract_text(response)

        except Exception as e:
            print(f"图片 PDF 提取失败: {e}")
            return ""

    def extract_text_from_docx(self, content: bytes) -> str:
        """从 DOCX 提取文本"""
        try:
            doc = Document(BytesIO(content))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"

            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text.strip()
        except Exception as e:
            print(f"DOCX 提取失败: {e}")
            raise ValueError(f"DOCX 文本提取失败: {str(e)}")

    def extract_text_from_doc(self, content: bytes, filename: str) -> str:
        """从旧版 .doc 文件提取文本"""
        import subprocess
        import tempfile

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            # 方法1: macOS 使用 textutil
            try:
                result = subprocess.run(
                    ['textutil', '-convert', 'txt', '-stdout', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            # 方法2: 尝试 antiword（如果安装了）
            try:
                result = subprocess.run(
                    ['antiword', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            # 方法3: 尝试 catdoc（如果安装了）
            try:
                result = subprocess.run(
                    ['catdoc', tmp_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout.strip()
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            # 如果都失败了，抛出错误
            raise ValueError("无法解析 .doc 文件。请将文件另存为 .docx 或 .pdf 格式后重试。")

        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass

    def quick_extract(self, text: str, filename: str) -> dict:
        """
        快速正则提取关键信息（不调用 AI，毫秒级）
        用于先入库，后续再异步精细解析
        """
        result = {
            "name": None,
            "phone": None,
            "email": None,
            "city": None,
            "current_company": None,
            "current_title": None,
            "years_of_experience": None,
            "expected_salary": None,
            "skills": [],
            "summary": None
        }

        # 只取文件名，去掉路径
        safe_filename = Path(filename).name if filename else ""

        # 1. 从文件名提取姓名
        if safe_filename:
            name_match = re.search(r'】([\u4e00-\u9fa5]{2,4})', safe_filename)
            if not name_match:
                name_match = re.match(r'^([\u4e00-\u9fa5]{2,4})[-_\s]', safe_filename.rsplit(".", 1)[0])
            if not name_match:
                name_match = re.search(r'([\u4e00-\u9fa5]{2,4})', safe_filename.rsplit(".", 1)[0])
            if name_match:
                result["name"] = name_match.group(1)

        # 2. 从文本提取姓名（如果文件名没提取到）
        if not result["name"]:
            # 常见格式：姓名：张三 / 姓名:张三 / Name: Zhang San
            name_patterns = [
                r'姓\s*名[：:]\s*([\u4e00-\u9fa5]{2,4})',
                r'^([\u4e00-\u9fa5]{2,4})\s*[\n\r]',  # 开头的中文名
            ]
            for pattern in name_patterns:
                match = re.search(pattern, text, re.MULTILINE)
                if match:
                    result["name"] = match.group(1)
                    break

        # 3. 提取手机号
        phone_match = re.search(r'1[3-9]\d{9}', text)
        if phone_match:
            result["phone"] = phone_match.group(0)

        # 4. 提取邮箱
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            result["email"] = email_match.group(0)

        # 5. 提取城市
        city_patterns = [
            r'(?:现居|所在城市|城市|工作地点|地点)[：:\s]*([\u4e00-\u9fa5]{2,4})',
            r'(北京|上海|广州|深圳|杭州|成都|武汉|南京|西安|苏州|天津|重庆)'
        ]
        for pattern in city_patterns:
            match = re.search(pattern, text)
            if match:
                result["city"] = match.group(1)
                break

        # 6. 提取工作年限
        exp_patterns = [
            r'(\d+)\s*[年+]\s*(?:工作)?经验',
            r'工作\s*(\d+)\s*年',
            r'(\d+)\s*years?\s*(?:of\s*)?experience',
        ]
        for pattern in exp_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result["years_of_experience"] = int(match.group(1))
                break

        # 7. 提取期望薪资
        salary_patterns = [
            r'期望[薪资薪酬月薪年薪][：:\s]*(\d+)[-~到]?(\d+)?[kKwW万]',
            r'(\d+)[-~到](\d+)[kKwW万]/[月年]',
        ]
        for pattern in salary_patterns:
            match = re.search(pattern, text)
            if match:
                try:
                    salary = int(match.group(1))
                    # 转换为万/年
                    if 'k' in pattern.lower() or 'K' in text[match.start():match.end()]:
                        salary = salary * 12 / 10  # k/月 -> 万/年
                    result["expected_salary"] = int(salary)
                except:
                    pass
                break

        # 8. 提取简短摘要（取前200字符作为摘要）
        clean_text = re.sub(r'\s+', ' ', text).strip()
        if len(clean_text) > 50:
            result["summary"] = clean_text[:200] + "..." if len(clean_text) > 200 else clean_text

        return result

    async def quick_parse(self, filename: str, content: bytes) -> dict:
        """
        快速解析模式：只提取文本 + 正则提取关键字段
        不调用 AI，速度快（毫秒级）
        """
        ext = filename.lower().split(".")[-1]
        text = ""
        extraction_method = "text"

        # 提取文本
        if ext == "pdf":
            text = self.extract_text_from_pdf(content)
            # 图片版 PDF 在快速模式下暂时跳过 AI 识别，标记为需要后续处理
            if len(text.strip()) < 50 and self.is_image_based_pdf(content):
                extraction_method = "pending_vision"
        elif ext == "docx":
            text = self.extract_text_from_docx(content)
        elif ext == "doc":
            text = self.extract_text_from_doc(content, filename)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

        # 保存文件
        file_path = await self.save_file(filename, content)

        # 快速正则提取
        quick_data = self.quick_extract(text, filename)

        return {
            "file_path": file_path,
            "file_name": filename,
            "file_type": ext,
            "raw_text": text,
            "parsed_data": quick_data,
            "extraction_method": extraction_method,
            "needs_ai_parsing": True  # 标记需要后续 AI 精细解析
        }

    async def parse(self, filename: str, content: bytes) -> dict:
        """完整解析模式（兼容原有逻辑）"""
        ext = filename.lower().split(".")[-1]
        text = ""
        extraction_method = "text"

        # 提取文本
        if ext == "pdf":
            # 先尝试普通文本提取
            text = self.extract_text_from_pdf(content)

            # 如果提取到的文本太少，尝试图片识别
            if len(text.strip()) < 50:
                print(f"PDF 文本提取结果较少 ({len(text)} 字符)，尝试图片识别...")
                if self.is_image_based_pdf(content):
                    extraction_method = "vision"
                    text = await self.extract_text_from_image_pdf(content)
                    if text:
                        print(f"图片识别成功，提取到 {len(text)} 字符")

        elif ext == "docx":
            text = self.extract_text_from_docx(content)
        elif ext == "doc":
            text = self.extract_text_from_doc(content, filename)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

        # 检查提取结果
        if not text or len(text.strip()) < 20:
            raise ValueError(f"无法从文件中提取足够的文本内容 (提取到 {len(text) if text else 0} 字符)")

        # 保存文件
        file_path = await self.save_file(filename, content)

        # 使用 AI 解析
        try:
            parsed = await self.ai.parse_resume(text)
        except Exception as e:
            print(f"AI 解析失败: {e}")
            parsed = {}

        return {
            "file_path": file_path,
            "file_name": filename,
            "file_type": ext,
            "raw_text": text,
            "parsed_data": parsed,
            "extraction_method": extraction_method
        }

    async def ai_parse_text(self, text: str) -> dict:
        """单独的 AI 解析方法，用于后台精细解析（优先使用 Gemini）"""
        try:
            return await self.ai.parse_resume_with_gemini(text)
        except Exception as e:
            print(f"AI 解析失败: {e}")
            return {}

    async def parse_batch(self, files: list[tuple[str, bytes]]) -> list[dict]:
        results = []
        for filename, content in files:
            try:
                result = await self.parse(filename, content)
                results.append({"success": True, **result})
            except Exception as e:
                results.append({"success": False, "filename": filename, "error": str(e)})
        return results
